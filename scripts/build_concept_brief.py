from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


OUT = Path("artifacts/LendSignal_Credit_Certificate_Brief.docx")


COLORS = {
    "ink": RGBColor(20, 31, 43),
    "muted": RGBColor(92, 103, 115),
    "blue": RGBColor(46, 116, 181),
    "dark_blue": RGBColor(31, 77, 120),
    "fill": "F2F4F7",
    "fill_blue": "E8EEF5",
    "border": "B7C2D0",
    "green": RGBColor(31, 95, 74),
}


def set_cell_shading(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), fill)
    tc_pr.append(shd)


def set_cell_border(cell, color="B7C2D0", size="6"):
    tc = cell._tc
    tc_pr = tc.get_or_add_tcPr()
    borders = tc_pr.first_child_found_in("w:tcBorders")
    if borders is None:
        borders = OxmlElement("w:tcBorders")
        tc_pr.append(borders)
    for edge in ("top", "left", "bottom", "right"):
        tag = "w:" + edge
        element = borders.find(qn(tag))
        if element is None:
            element = OxmlElement(tag)
            borders.append(element)
        element.set(qn("w:val"), "single")
        element.set(qn("w:sz"), size)
        element.set(qn("w:space"), "0")
        element.set(qn("w:color"), color)


def set_cell_margins(cell, top=80, start=120, bottom=80, end=120):
    tc = cell._tc
    tc_pr = tc.get_or_add_tcPr()
    margins = tc_pr.first_child_found_in("w:tcMar")
    if margins is None:
        margins = OxmlElement("w:tcMar")
        tc_pr.append(margins)
    for m, v in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
        node = margins.find(qn(f"w:{m}"))
        if node is None:
            node = OxmlElement(f"w:{m}")
            margins.append(node)
        node.set(qn("w:w"), str(v))
        node.set(qn("w:type"), "dxa")


def set_table_width(table, widths):
    table.autofit = False
    tbl_pr = table._tbl.tblPr
    tbl_w = tbl_pr.first_child_found_in("w:tblW")
    if tbl_w is None:
        tbl_w = OxmlElement("w:tblW")
        tbl_pr.append(tbl_w)
    tbl_w.set(qn("w:type"), "dxa")
    tbl_w.set(qn("w:w"), str(sum(widths)))

    indent = tbl_pr.first_child_found_in("w:tblInd")
    if indent is None:
        indent = OxmlElement("w:tblInd")
        tbl_pr.append(indent)
    indent.set(qn("w:type"), "dxa")
    indent.set(qn("w:w"), "120")

    grid = table._tbl.tblGrid
    if grid is None:
        grid = OxmlElement("w:tblGrid")
        table._tbl.insert(0, grid)
    for child in list(grid):
        grid.remove(child)
    for w in widths:
        col = OxmlElement("w:gridCol")
        col.set(qn("w:w"), str(w))
        grid.append(col)

    for row in table.rows:
        for idx, cell in enumerate(row.cells):
            tc_pr = cell._tc.get_or_add_tcPr()
            tc_w = tc_pr.first_child_found_in("w:tcW")
            if tc_w is None:
                tc_w = OxmlElement("w:tcW")
                tc_pr.append(tc_w)
            tc_w.set(qn("w:type"), "dxa")
            tc_w.set(qn("w:w"), str(widths[idx]))
            set_cell_margins(cell)
            set_cell_border(cell)
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER


def set_font(run, size=11, bold=False, color=None):
    run.font.name = "Calibri"
    run._element.rPr.rFonts.set(qn("w:ascii"), "Calibri")
    run._element.rPr.rFonts.set(qn("w:hAnsi"), "Calibri")
    run.font.size = Pt(size)
    run.bold = bold
    if color:
        run.font.color.rgb = color


def add_para(doc, text="", size=11, bold=False, color=None, after=6, before=0, align=None):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(before)
    p.paragraph_format.space_after = Pt(after)
    p.paragraph_format.line_spacing = 1.1
    if align is not None:
        p.alignment = align
    r = p.add_run(text)
    set_font(r, size=size, bold=bold, color=color or COLORS["ink"])
    return p


def add_heading(doc, text, level=1):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(14 if level == 1 else 8)
    p.paragraph_format.space_after = Pt(6)
    r = p.add_run(text)
    set_font(r, size=16 if level == 1 else 13, bold=True, color=COLORS["blue"] if level == 1 else COLORS["dark_blue"])
    return p


def add_bullet(doc, text, bold_prefix=None):
    p = doc.add_paragraph(style="List Bullet")
    p.paragraph_format.space_after = Pt(4)
    p.paragraph_format.line_spacing = 1.1
    if bold_prefix and text.startswith(bold_prefix):
        r1 = p.add_run(bold_prefix)
        set_font(r1, bold=True, color=COLORS["ink"])
        r2 = p.add_run(text[len(bold_prefix):])
        set_font(r2, color=COLORS["ink"])
    else:
        r = p.add_run(text)
        set_font(r, color=COLORS["ink"])
    return p


def add_callout(doc, title, body):
    table = doc.add_table(rows=1, cols=1)
    set_table_width(table, [9360])
    cell = table.cell(0, 0)
    set_cell_shading(cell, "EEF5F1")
    p = cell.paragraphs[0]
    p.paragraph_format.space_after = Pt(3)
    r = p.add_run(title)
    set_font(r, size=11, bold=True, color=COLORS["green"])
    p2 = cell.add_paragraph()
    p2.paragraph_format.space_after = Pt(0)
    r2 = p2.add_run(body)
    set_font(r2, size=10.5, color=COLORS["ink"])
    doc.add_paragraph().paragraph_format.space_after = Pt(2)


def add_kv_table(doc, rows):
    table = doc.add_table(rows=len(rows), cols=2)
    set_table_width(table, [2400, 6960])
    for i, (k, v) in enumerate(rows):
        c0, c1 = table.rows[i].cells
        set_cell_shading(c0, COLORS["fill"])
        p0 = c0.paragraphs[0]
        p0.paragraph_format.space_after = Pt(0)
        r0 = p0.add_run(k)
        set_font(r0, size=10, bold=True, color=COLORS["dark_blue"])
        p1 = c1.paragraphs[0]
        p1.paragraph_format.space_after = Pt(0)
        r1 = p1.add_run(v)
        set_font(r1, size=10, color=COLORS["ink"])
    doc.add_paragraph().paragraph_format.space_after = Pt(2)


def add_three_col_table(doc, headers, rows, compact=False):
    table = doc.add_table(rows=1 + len(rows), cols=3)
    set_table_width(table, [2100, 3860, 3400] if compact else [2500, 3430, 3430])
    for j, h in enumerate(headers):
        cell = table.rows[0].cells[j]
        set_cell_shading(cell, COLORS["fill_blue"])
        p = cell.paragraphs[0]
        p.paragraph_format.space_after = Pt(0)
        r = p.add_run(h)
        set_font(r, size=9.5 if compact else 10, bold=True, color=COLORS["dark_blue"])
    for i, row in enumerate(rows, start=1):
        for j, value in enumerate(row):
            cell = table.rows[i].cells[j]
            p = cell.paragraphs[0]
            p.paragraph_format.space_after = Pt(0)
            r = p.add_run(value)
            set_font(r, size=8.8 if compact else 9.5, color=COLORS["ink"], bold=(j == 0))
    doc.add_paragraph().paragraph_format.space_after = Pt(2)


def add_footer(section):
    footer = section.footer.paragraphs[0]
    footer.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    r = footer.add_run("LendSignal concept brief")
    set_font(r, size=9, color=COLORS["muted"])


def build():
    OUT.parent.mkdir(exist_ok=True)
    doc = Document()
    section = doc.sections[0]
    section.top_margin = Inches(0.72)
    section.bottom_margin = Inches(0.72)
    section.left_margin = Inches(0.82)
    section.right_margin = Inches(0.82)
    add_footer(section)

    styles = doc.styles
    styles["Normal"].font.name = "Calibri"
    styles["Normal"].font.size = Pt(11)

    add_para(doc, "Hackathon Concept Brief", size=10, bold=True, color=COLORS["muted"], after=4)
    title = add_para(doc, "LendSignal", size=27, bold=True, color=COLORS["ink"], after=2)
    title.alignment = WD_ALIGN_PARAGRAPH.LEFT
    add_para(
        doc,
        "Updateable onchain credit certificates for B2B working-capital lending",
        size=13,
        color=COLORS["dark_blue"],
        after=12,
    )
    add_callout(
        doc,
        "Name recommendation",
        "Use LendSignal for the hackathon: it is short, direct, and points to credit approval for business lending.",
    )

    add_heading(doc, "One-Liner")
    add_para(
        doc,
        "LendSignal turns a business wallet into an updateable onchain Credit Certificate, enabling working-capital loans through confidential verification, ENS-gated approval, and decentralized default protection.",
        size=11.5,
        bold=True,
        color=COLORS["ink"],
        after=8,
    )

    add_heading(doc, "Problem And Opportunity")
    add_bullet(doc, "Emerging market businesses need working capital, but onchain lending still mostly requires excessive collateral.")
    add_bullet(doc, "Protocols can see wallet activity, but not business documents, KYC/KYB evidence, repayment capacity, or fraud risk.")
    add_bullet(doc, "Confidential AI can validate private borrower evidence and output only a score, risk tier, attestation hash, and evidence digest.")
    add_bullet(doc, "A credit certificate lets lending vaults make repeatable credit decisions without exposing raw private data.")

    add_heading(doc, "Product Flow")
    add_three_col_table(
        doc,
        ["Step", "What Happens", "Public / Onchain Output"],
        [
            ("1. Onboard", "Connect wallet, enter company data, submit mock docs and KYC/KYB.", "Wallet becomes credit identity."),
            ("2. Verify", "Confidential AI Attester-style flow analyzes borrower evidence.", "AI score, risk tier, hashes."),
            ("3. Certify", "Combine AI, bureau, and wallet-behavior signals.", "Updateable Credit Certificate."),
            ("4. Lend", "Vault checks certificate, score, ENS gate, and liquidity.", "Automatic loan payout."),
            ("5. Protect", "Borrower fees fund a default pool for LPs.", "Lender reimbursement on default."),
        ],
        compact=True,
    )

    doc.add_page_break()
    add_heading(doc, "48-Hour MVP")
    add_kv_table(
        doc,
        [
            ("Business onboarding", "Wallet connect, business profile, mock document checklist, mock KYC/KYB status."),
            ("Credit score", "Confidential AI score 70% + bureau score 20% + wallet behavior 10%."),
            ("Credit certificate", "Registry stores score, risk tier, status, attestation hash, evidence digest, expiration."),
            ("Lending vault", "Approves and pays out when certificate is active, score >= 750, ENS gate passed, and liquidity exists."),
            ("Default fund", "LP deposits plus borrower fees; demo reimbursement on default."),
        ],
    )

    add_heading(doc, "Credit Certificate")
    add_para(doc, "The Credit Certificate is the public representation of the business credit state. It is updateable over time as new documents, bureau signals, wallet behavior, repayments, or defaults are observed.")
    add_kv_table(
        doc,
        [
            ("Expose", "business wallet, certificate id, combined score, risk tier, status, attestation hash, evidence digest, ENS gate, issued/expiry dates."),
            ("Hide", "raw documents, bank statements, full KYC/KYB records, tax records, private invoices, personal identity documents."),
            ("Lifecycle", "pending, active, updated, expired, revoked, defaulted."),
            ("Vault policy", "approve if active, not expired, score >= 750, ENS gate passed, and vault has liquidity."),
        ],
    )

    add_heading(doc, "Chainlink Confidential AI Attester Integration")
    add_bullet(doc, "Base URL: https://confidential-ai-dev-preview.cldev.cloud")
    add_bullet(doc, "Flow: POST /v1/inference, then poll GET /v1/inference/:id until completed or failed.")
    add_bullet(doc, "Resources: uploaded files or HTTP(S) URLs, up to 10 resources, each capped at 10 MiB.")
    add_bullet(doc, "Models: gemma4 or qwen3.6.")
    add_bullet(doc, "Hackathon fallback: if no key is available, use deterministic mock outputs with the same response shape.")

    doc.add_page_break()
    add_heading(doc, "Minimal Architecture")
    add_three_col_table(
        doc,
        ["Component", "Build For Hackathon", "Why It Matters"],
        [
            ("Frontend", "Onboarding, score, certificate, vault, default fund screens.", "Makes the borrower and lender story obvious."),
            ("Score service", "Mock or real attester API call, bureau mock, wallet behavior mock.", "Creates the credit decision input."),
            ("Certificate registry", "Issue, update, revoke, read certificate.", "Turns credit state into reusable onchain primitive."),
            ("Lending vault", "Deposit, request loan, approve and payout.", "Demonstrates non-collateral-only lending."),
            ("Default fund", "LP deposits, borrower fees, reimburse default.", "Explains lender protection and business model."),
        ],
    )

    add_heading(doc, "Business Model")
    add_bullet(doc, "Borrower certification fee: business pays to create or refresh a certificate.")
    add_bullet(doc, "Loan origination fee: fee is charged when a vault pays out a loan.")
    add_bullet(doc, "Default fund fee: part of borrower fees funds LP yield and lender protection.")
    add_bullet(doc, "Protocol integration fee: lending pools pay for certificate checks, webhooks, or API access.")
    add_bullet(doc, "Roadmap: agent-based monitoring, ERC-8004-style reputation, curator marketplace, real bureau integrations.")

    add_heading(doc, "Pitch Close")
    add_para(
        doc,
        "LendSignal gives emerging market businesses a portable credit identity and gives lenders a safer way to approve undercollateralized working-capital loans.",
        size=12,
        bold=True,
        color=COLORS["green"],
        after=4,
    )

    doc.save(OUT)


if __name__ == "__main__":
    build()

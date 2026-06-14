from pathlib import Path

from docx import Document
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


OUT = Path("artifacts/LendSignal_User_Data_Flow.docx")

INK = RGBColor(20, 31, 43)
MUTED = RGBColor(92, 103, 115)
BLUE = RGBColor(46, 116, 181)
DARK_BLUE = RGBColor(31, 77, 120)
GREEN = RGBColor(31, 95, 74)


def set_font(run, size=10.5, bold=False, color=INK):
    run.font.name = "Calibri"
    run._element.rPr.rFonts.set(qn("w:ascii"), "Calibri")
    run._element.rPr.rFonts.set(qn("w:hAnsi"), "Calibri")
    run.font.size = Pt(size)
    run.bold = bold
    run.font.color.rgb = color


def para(doc, text="", size=10.5, bold=False, color=INK, before=0, after=5, align=None):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(before)
    p.paragraph_format.space_after = Pt(after)
    p.paragraph_format.line_spacing = 1.08
    if align is not None:
        p.alignment = align
    r = p.add_run(text)
    set_font(r, size=size, bold=bold, color=color)
    return p


def heading(doc, text):
    return para(doc, text, size=14, bold=True, color=BLUE, before=8, after=3)


def bullet(doc, text):
    p = doc.add_paragraph(style="List Bullet")
    p.paragraph_format.space_after = Pt(2)
    p.paragraph_format.line_spacing = 1.05
    r = p.add_run(text)
    set_font(r, size=10, color=INK)


def shade(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), fill)
    tc_pr.append(shd)


def borders(cell):
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_borders = tc_pr.first_child_found_in("w:tcBorders")
    if tc_borders is None:
        tc_borders = OxmlElement("w:tcBorders")
        tc_pr.append(tc_borders)
    for edge in ("top", "left", "bottom", "right"):
        node = tc_borders.find(qn(f"w:{edge}"))
        if node is None:
            node = OxmlElement(f"w:{edge}")
            tc_borders.append(node)
        node.set(qn("w:val"), "single")
        node.set(qn("w:sz"), "6")
        node.set(qn("w:space"), "0")
        node.set(qn("w:color"), "B7C2D0")


def margins(cell):
    tc_pr = cell._tc.get_or_add_tcPr()
    mar = tc_pr.first_child_found_in("w:tcMar")
    if mar is None:
        mar = OxmlElement("w:tcMar")
        tc_pr.append(mar)
    for key, val in {"top": 55, "bottom": 55, "start": 105, "end": 105}.items():
        node = mar.find(qn(f"w:{key}"))
        if node is None:
            node = OxmlElement(f"w:{key}")
            mar.append(node)
        node.set(qn("w:w"), str(val))
        node.set(qn("w:type"), "dxa")


def set_widths(table, widths):
    table.autofit = False
    tbl_pr = table._tbl.tblPr
    tbl_w = tbl_pr.first_child_found_in("w:tblW")
    if tbl_w is None:
        tbl_w = OxmlElement("w:tblW")
        tbl_pr.append(tbl_w)
    tbl_w.set(qn("w:type"), "dxa")
    tbl_w.set(qn("w:w"), str(sum(widths)))

    grid = table._tbl.tblGrid
    if grid is None:
        grid = OxmlElement("w:tblGrid")
        table._tbl.insert(0, grid)
    for child in list(grid):
        grid.remove(child)
    for width in widths:
        col = OxmlElement("w:gridCol")
        col.set(qn("w:w"), str(width))
        grid.append(col)

    for row in table.rows:
        for i, cell in enumerate(row.cells):
            tc_pr = cell._tc.get_or_add_tcPr()
            tc_w = tc_pr.first_child_found_in("w:tcW")
            if tc_w is None:
                tc_w = OxmlElement("w:tcW")
                tc_pr.append(tc_w)
            tc_w.set(qn("w:type"), "dxa")
            tc_w.set(qn("w:w"), str(widths[i]))
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            margins(cell)
            borders(cell)


def table(doc, headers, rows, widths):
    t = doc.add_table(rows=1 + len(rows), cols=len(headers))
    set_widths(t, widths)
    for i, header in enumerate(headers):
        cell = t.rows[0].cells[i]
        shade(cell, "E8EEF5")
        p = cell.paragraphs[0]
        p.paragraph_format.space_after = Pt(0)
        r = p.add_run(header)
        set_font(r, size=9, bold=True, color=DARK_BLUE)
    for r_idx, row in enumerate(rows, start=1):
        for c_idx, value in enumerate(row):
            cell = t.rows[r_idx].cells[c_idx]
            if c_idx == 0:
                shade(cell, "F2F4F7")
            p = cell.paragraphs[0]
            p.paragraph_format.space_after = Pt(0)
            run = p.add_run(value)
            set_font(run, size=8.8, bold=(c_idx == 0), color=INK)
    para(doc, "", after=0)


def build():
    OUT.parent.mkdir(exist_ok=True)
    doc = Document()
    section = doc.sections[0]
    section.top_margin = Inches(0.62)
    section.bottom_margin = Inches(0.62)
    section.left_margin = Inches(0.82)
    section.right_margin = Inches(0.82)

    footer = section.footer.paragraphs[0]
    footer.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    r = footer.add_run("LendSignal user data flow")
    set_font(r, size=9, color=MUTED)

    para(doc, "Operational Flow", size=10, bold=True, color=MUTED, after=3)
    para(doc, "LendSignal User Data Flow", size=25, bold=True, color=INK, after=2)
    para(doc, "From business onboarding to an ENS-indexed Credit Certificate", size=12.5, color=DARK_BLUE, after=12)

    para(
        doc,
        "The user submits business data, documents and wallet identity. LendSignal processes that information offchain, generates a credit score, updates the CreditCertificateRegistry, and indexes the certificate through ENS records.",
        size=11.2,
        bold=True,
        after=10,
    )

    heading(doc, "1. What The User Fills In")
    table(
        doc,
        ["Area", "Fields", "MVP Handling"],
        [
            ("Wallet", "business wallet, network", "Connected wallet becomes credit identity."),
            ("Business profile", "legal name, country, industry, revenue band, requested loan", "Saved in app state/mock backend."),
            ("Documents", "registration, tax record, statements, invoices, signer auth", "Mock upload/checklist; raw files stay offchain."),
            ("KYC/KYB", "business verified, signer verified", "Mock passed/pending/failed status."),
            ("ENS", "business ENS name", "Resolved and later used for certificate discovery."),
        ],
        [2100, 4360, 2900],
    )

    heading(doc, "2. Processing Pipeline")
    table(
        doc,
        ["Layer", "Input", "Output"],
        [
            ("CRS Adapter", "business profile + principal metadata", "bureau score, payment risk, fraud risk, CRS report hash"),
            ("Chainlink AI", "documents + profile + prompt", "AI score, risk tier, attestation hash, evidence digest"),
            ("Wallet Analyzer", "wallet activity", "wallet behavior score and risk flags"),
            ("Score Combiner", "AI 60% + CRS 25% + wallet 15%", "combined score and final risk tier"),
        ],
        [2100, 3650, 3610],
    )

    doc.add_page_break()
    heading(doc, "3. Certificate Update")
    para(doc, "CreditCertificateRegistry receives only public-safe processed information.", size=10.5)
    table(
        doc,
        ["Stored Onchain", "Never Stored Onchain"],
        [
            ("business wallet", "raw PDFs"),
            ("combined score and risk tier", "full CRS report"),
            ("certificate status and expiration", "bank statements or tax records"),
            ("attestation hash, CRS hash, evidence digest", "KYC/KYB documents or personal identity files"),
        ],
        [4680, 4680],
    )

    heading(doc, "4. ENS Indexing")
    para(doc, "ENS is the discovery layer. CreditCertificateRegistry remains the source of truth.", size=10.5, bold=True)
    table(
        doc,
        ["ENS Text Record", "Value"],
        [
            ("lendsignal.registry", "<CreditCertificateRegistry address>"),
            ("lendsignal.certificate", "<certificate id>"),
            ("lendsignal.attestation", "<Chainlink attestation hash>"),
            ("lendsignal.crs-report", "<CRS report hash>"),
            ("lendsignal.risk-tier", "<risk tier>"),
            ("lendsignal.agent", "<LendSignal agent ENS name>"),
        ],
        [3100, 6260],
    )

    heading(doc, "5. Lending Decision")
    table(
        doc,
        ["Check", "Pass Condition"],
        [
            ("ENS resolution + records", "Name resolves to borrower wallet and points to registry/certificate."),
            ("Certificate", "Active and not expired."),
            ("Score", "Combined score above lending threshold."),
            ("Risk tier", "Low or medium default risk."),
            ("Liquidity", "Vault has enough available funds."),
        ],
        [2800, 6560],
    )

    heading(doc, "6. Update Triggers")
    para(
        doc,
        "New documents, refreshed CRS report, re-run Chainlink AI review, wallet behavior changes, repayment/default events, or KYC/KYB status changes.",
        size=9.8,
        after=3,
    )

    heading(doc, "Hackathon Version")
    para(
        doc,
        "Use mocked documents, mocked KYC/KYB, CRS_USE_MOCK=true, and real or mocked Chainlink AI depending on API key availability. The demo should still update CreditCertificateRegistry and show ENS records pointing to the active certificate.",
        size=10,
        bold=True,
        color=GREEN,
    )

    doc.save(OUT)


if __name__ == "__main__":
    build()
